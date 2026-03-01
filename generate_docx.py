"""
Generates blog-post.docx for the Power Platform Blog Digest article.
Run: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ────────────────────────────────────────────────────────────────

def set_run_color(run, hex_color):
    r, g, b = int(hex_color[0:2],16), int(hex_color[2:4],16), int(hex_color[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)

def add_heading(doc, text, level=2, color="0078d4"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_color(run, color)
        run.font.bold = True
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_callout(doc, title, text, color="0078d4"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF5FF')
    pPr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = p.add_run(title.upper() + "  ")
    r1.bold = True
    set_run_color(r1, color)
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_screenshot(doc, image_path, caption, series_title=None):
    """Embed an actual screenshot with optional series label and caption."""
    if series_title:
        t = doc.add_paragraph()
        t.paragraph_format.space_before = Pt(12)
        t.paragraph_format.space_after = Pt(4)
        r = t.add_run(series_title.upper())
        r.bold = True
        r.font.size = Pt(9)
        set_run_color(r, "0078d4")

    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        run.add_picture(image_path, width=Inches(5.8))
    else:
        # Fallback if image not found
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[ IMAGE NOT FOUND: {image_path} ]")
        r.bold = True
        set_run_color(r, "ca5010")
        r.font.size = Pt(11)

    cap = doc.add_paragraph(caption)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    for run in cap.runs:
        run.font.italic = True
        run.font.size = Pt(10)
        set_run_color(run, "666666")

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '0078D4')
        tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = table.rows[ri+1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            if ri % 2 == 1:
                tcPr = row.cells[ci]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F4F8FD')
                tcPr.append(shd)
    doc.add_paragraph()  # spacer

# Screenshot paths (relative to script location)
SS = "Screenshots"
IMG = {
    "email1":    f"{SS}/Sample digest email as it appears in Outlook \u2014 showing the branded header, post cards, and footer - Chunk1.png",
    "email2":    f"{SS}/Sample digest email as it appears in Outlook \u2014 showing the branded header, post cards, and footer - Chunk2.png",
    "email3":    f"{SS}/Sample digest email as it appears in Outlook \u2014 showing the branded header, post cards, and footer - Chunk3.png",
    "teams_full":f"{SS}/Sample digest email as it appears in Outlook \u2014 showing the branded header, post cards, and footer - Teams Notifications (extra boinus screenshot).png",
    "flow":      f"{SS}/Power Automate designer view showing the full flow \u2014 trigger, Main Scope, and Error Handler Scope.png",
    "envvar1":   f"{SS}/Flow parameters pane showing NotificationEmail and AlertEmail configuration fields - chunk1.png",
    "envvar2":   f"{SS}/Flow parameters pane showing NotificationEmail and AlertEmail configuration fields - chunk2.png",
    "card_email":f"{SS}/Close-up of a single post card in the digest email showing title, date, summary, and Read More button.png",
    "card_teams":f"{SS}/Close-up of a single post card in the digest email showing title, date, summary, and Read More button - optional teams card closeup view.png",
    "error":     f"{SS}/Error alert email generated by the Error Handler Scope \u2014 showing error code, message, and timestamp.png",
    "sol_all":   f"{SS}/Power Platform Solutions \u2014 Flow & connection references for RSS and Office 365 Outlook connectors.png",
    "sol_conn":  f"{SS}/Power Platform Solutions \u2014 connection references for RSS and Office 365 Outlook connectors.png",
}

# ── Build the document ──────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Title block ──────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "Never Miss a Power Platform Update:\n"
    "Building a Daily Blog Digest with Power Automate"
)
title_run.bold = True
title_run.font.size = Pt(26)
set_run_color(title_run, "0078d4")
title_p.paragraph_format.space_after = Pt(6)

# subtitle
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_r = sub_p.add_run(
    "How I automated a personalized HTML email digest and Microsoft Teams channel notification\n"
    "for the Power Platform blog — using RSS, Office 365 Outlook, Microsoft Teams,\n"
    "and Power Automate expressions — all in under an hour."
)
sub_r.font.size = Pt(13)
set_run_color(sub_r, "444444")
sub_r.font.italic = True
sub_p.paragraph_format.space_after = Pt(4)

# author / meta line
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_r = meta_p.add_run(
    "Ranjith Neela  \u00b7  Power Platform Architect  \u00b7  EPMPoint  \u00b7  February 28, 2026\n"
    "ranjith.neela@epmpoint.com"
)
meta_r.font.size = Pt(11)
set_run_color(meta_r, "0078d4")
meta_p.paragraph_format.space_after = Pt(24)

# divider
hr = doc.add_paragraph("\u2500" * 80)
hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr.paragraph_format.space_after = Pt(20)

# ── Section 1: The Problem ────────────────────────────────────────────────────
add_heading(doc, "The Problem: Information Overload", 1, "0078d4")
add_body(doc,
    "If you work in the Microsoft ecosystem \u2014 and especially with Power Platform \u2014 you already "
    "know the challenge: the official Microsoft Power Platform Blog publishes new content "
    "constantly. Release announcements, feature previews, admin tips, community spotlights \u2014 "
    "keeping up means bookmarking the site, setting browser alerts, or relying on social media "
    "feeds that bury the signal in noise."
)
add_body(doc,
    "I wanted something cleaner: a single daily notification that shows exactly what was published "
    "in the last 24 hours, formatted beautifully, with one-click links to each article \u2014 "
    "delivered to my inbox and to my team's Microsoft Teams channel. And I wanted it to be "
    "zero-maintenance \u2014 no servers, no code to deploy, no subscriptions."
)
add_body(doc, "The answer, naturally, was Power Automate.")

add_callout(doc,
    "\U0001f4a1 What You Will Learn",
    "In this post I walk through the design, the flow logic, and the key decisions behind "
    "the Power Platform Blog Digest solution \u2014 a fully exportable Power Automate cloud flow "
    "that monitors the Microsoft Power Platform RSS feed and delivers a polished HTML digest "
    "email and a Microsoft Teams Adaptive Card notification every morning."
)

# ── Section 2: What it does ──────────────────────────────────────────────────
add_heading(doc, "What the Flow Does \u2014 At a Glance", 1, "0078d4")
add_body(doc,
    "The flow is simple by design. Every day at 8:00 AM UTC it wakes up, checks the Power "
    "Platform blog's public RSS feed for posts published in the previous 24 hours, and either:"
)
add_bullet(doc,
    " a styled HTML digest email via Office 365 Outlook and posts an Adaptive Card to "
    "Microsoft Teams if new posts are found, or",
    "Sends"
)
add_bullet(doc, " gracefully (status: Succeeded, no notifications sent) if there is nothing new.", "Exits")
add_body(doc,
    "A separate error-handling scope catches any unexpected failures and fires an alert email "
    "to an admin address \u2014 so you always know if something went wrong."
)

# Email screenshots series
add_screenshot(doc, IMG["email1"],
    "The email arrives with subject 'New Microsoft Power Platform Blog Posts \u2014 [date]' \u2014 the branded header and first post card are immediately visible",
    series_title="Sample digest email as received in Microsoft Outlook (1 of 3)"
)
add_screenshot(doc, IMG["email2"],
    "Scrolling reveals additional post cards \u2014 each with title, publish date, summary, and a 'Read More' button linking directly to the article",
    series_title="Sample digest email \u2014 continued (2 of 3)"
)
add_screenshot(doc, IMG["email3"],
    "The email closes with the final post card and a subtle footer with a 'View Full Blog Archive' link",
    series_title="Sample digest email \u2014 footer (3 of 3)"
)

# ── Section 3: Flow Architecture ─────────────────────────────────────────────
add_heading(doc, "Flow Architecture \u2014 Step by Step", 1, "0078d4")
add_body(doc,
    "The flow has three top-level building blocks: an initialization step, a Main Scope "
    "that does all the real work, and an Error Handler Scope that runs only on failure."
)

steps = [
    ("1. Trigger \u2014 Recurrence (Daily at 8:00 AM UTC)",
     "A simple recurrence trigger fires once per day. No event, no webhook \u2014 just a reliable "
     "schedule so the flow is predictable and easy to monitor in run history."),
    ("2. Initialize Variables \u2014 HTML_Body",
     "Empty string variables are declared before the main scope. The primary accumulator "
     "(HTML_Body) is filled with styled post cards during the loop and later injected between "
     "the email header and footer."),
    ("3. List All RSS Feed Items & Filter (last 24 hours)",
     "The built-in RSS connector calls the Power Platform blog feed. A Filter Array action "
     "limits results to items published within the last 24 hours using addDays(utcNow(), -1)."),
    ("4. Condition \u2014 Any New Posts?",
     "A condition checks length(body('List_all_RSS_feed_items')) > 0. If false, a Terminate "
     "action marks the run Succeeded with a friendly message and no notifications are sent."),
    ("5. Apply to Each \u2014 Build HTML Cards",
     "For every RSS item, a Compose action renders a self-contained styled div card containing "
     "the post title (linked), publish date, summary text, and a 'Read More' button. Each card "
     "is then appended to the HTML_Body variable."),
    ("6. Compose Email Header & Footer",
     "After the loop, two Compose actions generate a branded gradient header banner (with "
     "today's date) and a footer with a link to the full blog archive."),
    ("7. Compose Full Email Body & Final Card",
     "A single expression merges Header + HTML_Body variable + Footer into the final email body "
     "string. A separate Compose assembles the Teams Adaptive Card payload from the same post data."),
    ("8. Send Notification Email (Office 365 Outlook)",
     "The HTML digest is dispatched via the Office 365 Outlook connector to the NotificationEmail "
     "environment variable recipient. The subject line includes today's date for easy identification."),
    ("T. Post Card in a Chat or Channel (Microsoft Teams)",
     "Immediately after the email, the flow posts the digest as an Adaptive Card to a configured "
     "Microsoft Teams chat or channel \u2014 giving the whole team instant visibility without anyone "
     "needing to check their inbox."),
    ("! Error Handler Scope (runs only on failure)",
     "If Main Scope fails or times out, a second scope builds an HTML error report \u2014 capturing "
     "the error code and message \u2014 and sends it to a separate AlertEmail environment variable."),
]

for title, body in steps:
    p = doc.add_paragraph(style='List Number')
    r1 = p.add_run(title + "\n")
    r1.bold = True
    r1.font.size = Pt(12)
    set_run_color(r1, "0078d4")
    r2 = p.add_run(body)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)

add_screenshot(doc, IMG["flow"],
    "Power Automate designer view \u2014 the complete flow from Recurrence trigger through Main Scope "
    "(including the 'Post card in a chat or channel' Teams step) to the Error Handler Scope",
    series_title="Power Automate designer \u2014 full flow view"
)

# ── Section 4: Design Decisions ───────────────────────────────────────────────
add_heading(doc, "Key Design Decisions", 1, "0078d4")

add_heading(doc, "Why Two Scopes?", 2, "005a9e")
add_body(doc,
    "Wrapping the main logic in a Scope action is the cleanest way to implement try/catch "
    "semantics in Power Automate. The Error Handler Scope is configured with "
    "runAfter: { Main_Scope: [\"Failed\", \"TimedOut\"] }, which means it is completely "
    "invisible during a healthy run and only activates when something goes wrong. This pattern "
    "keeps the happy-path logic clean and the error-path logic isolated."
)

add_heading(doc, "Graceful Exit When No Posts Are Found", 2, "005a9e")
add_body(doc,
    "Rather than sending empty notifications on quiet days, the flow uses a Terminate action "
    "inside the else branch with runStatus: \"Succeeded\". This keeps your run history green "
    "(not red) for no-news days, which matters when you are monitoring many flows."
)

add_heading(doc, "Dual-Channel Delivery: Email and Teams", 2, "005a9e")
add_body(doc,
    "The flow delivers the same digest content through two channels in a single run. The "
    "Office 365 Outlook action sends a rich, self-contained HTML email for in-depth reading, "
    "while the Teams connector posts an Adaptive Card to a shared channel \u2014 giving the whole "
    "team at-a-glance visibility without anyone needing to check their inbox. Both notifications "
    "fire from the same flow run, so there is zero duplication of effort."
)

add_heading(doc, "HTML Email Built Entirely in Flow Expressions", 2, "005a9e")
add_body(doc,
    "Every piece of HTML \u2014 from the gradient header to the individual post cards \u2014 is generated "
    "using Power Automate Compose actions and inline expressions. There is no external template "
    "engine, no Azure Function, no storage account. The entire rendering pipeline lives inside "
    "the flow definition."
)

add_heading(doc, "Environment Variables Instead of Hard-Coded Values", 2, "005a9e")
add_body(doc,
    "The NotificationEmail and AlertEmail values are defined as environment variables within the "
    "Power Platform solution \u2014 not hard-coded inside actions. This means anyone who imports the "
    "solution can set the email addresses directly in the solution's Environment Variables section "
    "without touching the flow internals, making it truly plug-and-play across environments."
)

add_screenshot(doc, IMG["envvar1"],
    "The solution's Environment Variables section \u2014 AlertEmail and NotificationEmail are configured "
    "here once, keeping the flow portable across environments without any internal edits",
    series_title="Environment variables in the solution (1 of 2)"
)
add_screenshot(doc, IMG["envvar2"],
    "Inside the flow \u2014 the Send Notification Email action resolves the NotificationEmail environment "
    "variable at runtime; the Error Handler Scope below uses the same pattern for AlertEmail",
    series_title="Environment variables used inside the flow (2 of 2)"
)

# ── Section 5: Notification Output ────────────────────────────────────────────
add_heading(doc, "What the Notifications Look Like", 1, "0078d4")

add_heading(doc, "Email Digest in Outlook", 2, "005a9e")
add_body(doc,
    "The generated email is fully self-contained HTML \u2014 no external CSS, no linked stylesheets "
    "\u2014 so it renders reliably across email clients."
)

add_table(doc,
    ["Element", "Design Choice"],
    [
        ["Header",      "Microsoft blue gradient banner with today's date \u2014 instantly recognisable"],
        ["Post Cards",  "White card with blue left border, title in brand blue, publish date, summary, and a 'Read More' CTA button"],
        ["Footer",      "Subtle grey separator with a direct link to the full blog archive"],
        ["Font",        "Segoe UI \u2014 consistent with Microsoft product design"],
        ["Max width",   "720 px \u2014 optimal for desktop email clients; readable on mobile"],
    ]
)

add_screenshot(doc, IMG["card_email"],
    "A single email post card \u2014 article title linked to the full post, calendar-icon publish date, "
    "rich summary text, and a prominent 'Read More' button",
    series_title="Email post card close-up"
)

add_heading(doc, "Microsoft Teams Channel Card", 2, "005a9e")
add_body(doc,
    "In parallel with the email, the flow posts an Adaptive Card from the EPMPoint Workflows bot "
    "to the configured Teams chat or channel. The card presents each new post with its title, "
    "publish date, summary, and a 'Read More' button \u2014 giving the team instant awareness of "
    "the day's Power Platform news directly in their Teams feed, without switching to email."
)

add_screenshot(doc, IMG["card_teams"],
    "Close-up of the Teams Adaptive Card \u2014 the 'EPMPoint Intelligence Digest' card lists each new "
    "post with title, publish date, summary, and a Read More button, mirroring the email digest inside Teams",
    series_title="Teams Adaptive Card close-up (1 of 2)"
)
add_screenshot(doc, IMG["teams_full"],
    "Full Teams view \u2014 the card appears in the Workflows chat showing all four new blog posts, "
    "with EPMPoint branding and 'View Full Blog Archive' button at the bottom",
    series_title="Teams full notification view (2 of 2)"
)

add_screenshot(doc, IMG["error"],
    "Error alert email from the Error Handler Scope \u2014 sent as high importance, showing flow name, "
    "UTC timestamp, error code, and error message; prompts the admin to review run history in the portal",
    series_title="Error Handler Scope \u2014 alert email"
)

# ── Section 6: Prerequisites & Import ─────────────────────────────────────────
add_heading(doc, "Prerequisites & How to Import", 1, "0078d4")

add_callout(doc,
    "\U0001f4cb What You Need",
    "A Power Platform environment with Dataverse  \u00b7  An Office 365 Outlook licence  \u00b7  "
    "A Microsoft Teams licence  \u00b7  Maker permissions in the target environment",
    "ca5010"
)

steps_import = [
    "Download the solution ZIP \u2014 PowerPlatformBlogDigest_1_0_0_3.zip",
    "Navigate to make.powerautomate.com \u2192 your environment \u2192 Solutions",
    "Click Import solution and upload the ZIP",
    "During import, create or select connections for RSS, Office 365 Outlook, and Microsoft Teams",
    "After import, open the solution and set values for the NotificationEmail and AlertEmail environment variables",
    "Open the flow and configure the 'Post card in a chat or channel' action with your target Teams team and channel",
    "Save and Turn on the flow",
]
for s in steps_import:
    add_numbered(doc, s)

add_screenshot(doc, IMG["sol_all"],
    "The imported solution contains 4 objects \u2014 one Cloud Flow ('Power Platform Blog Digest') and "
    "three Connection References for Microsoft Teams, Office 365 Outlook, and RSS",
    series_title="Power Platform solution contents after import (1 of 2)"
)
add_screenshot(doc, IMG["sol_conn"],
    "Connection References detail \u2014 all three connectors (Microsoft Teams, Office 365 Outlook, RSS) "
    "must be linked to active connections before the flow will run",
    series_title="Solution connection references (2 of 2)"
)

# ── Section 7: Extending ──────────────────────────────────────────────────────
add_heading(doc, "Ideas for Extending This Flow", 1, "0078d4")
add_body(doc,
    "The current flow delivers a daily dual-channel digest with full error handling and does "
    "it well. Natural extensions include:"
)
add_bullet(doc, " \u2014 Add the Power Apps, Power BI, and Copilot Studio blogs to the same digest by parallelising additional RSS actions", "Multiple RSS feeds")
add_bullet(doc, " \u2014 Change the recurrence to run on Mondays with a -7 day offset to capture the full previous week in one card", "Weekly digest mode")
add_bullet(doc, " \u2014 Filter specifically for IT Pro or Developer posts using the blog's audience query parameters", "Audience filtering")
add_bullet(doc, " \u2014 Write each digest run to a Dataverse table for audit and trend analysis", "Dataverse logging")
add_bullet(doc, " \u2014 Flag posts matching keywords like 'Copilot', 'deprecation', or 'breaking change'", "Keyword alerts")

add_callout(doc,
    "\u2705 Pro Tip",
    "If you want the digest to cover the weekend, consider changing the recurrence to run on "
    "Mondays with a -3 day offset instead of -1. That way Friday, Saturday, and Sunday posts "
    "all appear in the Monday morning email and Teams card.",
    "107c10"
)

# ── Section 8: Closing ────────────────────────────────────────────────────────
add_heading(doc, "Wrapping Up", 1, "0078d4")
add_body(doc,
    "The Power Platform Blog Digest is a great example of what I love about Power Automate: "
    "you can solve a real, everyday problem without writing a single line of server-side code, "
    "without provisioning any infrastructure, and without any ongoing maintenance burden. "
    "The entire solution \u2014 from trigger to dual-channel notification to error handling \u2014 lives "
    "inside a single exported solution ZIP that you can import into any environment in minutes."
)
add_body(doc,
    "Whether your team prefers email or Teams (or both), this flow has you covered. You get a "
    "rich HTML Outlook digest for in-depth reading and a quick Teams Adaptive Card for at-a-glance "
    "awareness \u2014 all from the same single flow run, zero extra effort."
)
add_body(doc,
    "The solution is packaged as PowerPlatformBlogDigest_1_0_0_3.zip and is ready to import "
    "into any environment today. If you find this useful, feel free to extend it, remix it, "
    "and share it with your team."
)

# ── Author Bio ────────────────────────────────────────────────────────────────
doc.add_paragraph()
hr2 = doc.add_paragraph("\u2500" * 80)
hr2.paragraph_format.space_before = Pt(16)

bio_heading = doc.add_paragraph()
bio_r = bio_heading.add_run("About the Author")
bio_r.bold = True
bio_r.font.size = Pt(14)
set_run_color(bio_r, "0078d4")
bio_heading.paragraph_format.space_after = Pt(6)

bio_name = doc.add_paragraph()
name_r = bio_name.add_run("Ranjith Neela")
name_r.bold = True
name_r.font.size = Pt(13)
set_run_color(name_r, "0078d4")

bio_role = doc.add_paragraph()
role_r = bio_role.add_run("Power Platform Architect  |  EPMPoint  |  ranjith.neela@epmpoint.com")
role_r.font.size = Pt(11)
set_run_color(role_r, "666666")
role_r.font.italic = True
bio_role.paragraph_format.space_after = Pt(8)

add_body(doc,
    "Ranjith is a Power Platform architect at EPMPoint, specialising in Power Automate, "
    "Power Apps, and Dataverse solutions for enterprise clients. He is passionate about "
    "building elegant, low-code automations that eliminate repetitive work and surface "
    "the right information at the right time."
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "blog-post.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
